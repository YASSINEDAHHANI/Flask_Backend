import os
from datetime import datetime, timezone
from functools import wraps
import uuid
import json
from bson import ObjectId
import langdetect
import anthropic
import PyPDF2
import docx
from flask import Flask, request, jsonify, session, Response
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from pymongo import MongoClient
from dotenv import load_dotenv
from cryptography.fernet import Fernet
import base64
from admin import admin_bp, init_admin_collections
from manager import manager_bp
import re
import io
from io import BytesIO
from flask import send_file
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import admin
import manager
from datetime import datetime, timezone
import traceback
import anthropic
import json
from datetime import datetime, timezone
import os
from flask_cors import CORS, cross_origin
from minio import Minio
from minio.error import S3Error
from flask_cors import CORS
try:
    from minio_rag_system import MinIORAGSystem
    MINIO_RAG_AVAILABLE = True
    print("✅ MinIO RAG System imported successfully")
except ImportError as e:
    MINIO_RAG_AVAILABLE = False
    print(f"❌ CRITICAL ERROR: MinIO RAG System not available: {e}")
    print("❌ Application cannot start without MinIO RAG System")

load_dotenv()

class MongoJSONEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, ObjectId):
            return str(obj)
        if isinstance(obj, datetime):
            return obj.isoformat()
        return super().default(obj)

app = Flask(__name__)
app.register_blueprint(admin_bp)
app.register_blueprint(manager_bp)
app.json_encoder = MongoJSONEncoder
app.secret_key = os.getenv("SECRET_KEY", "supersecret")

app.config.update(
    SESSION_COOKIE_NAME="flask_session",
    SESSION_COOKIE_SECURE=False,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    PERMANENT_SESSION_LIFETIME=86400,
    SESSION_REFRESH_EACH_REQUEST=True
)


cors = CORS(app, supports_credentials=True)
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["300 per minute"] 
)

MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "localhost:9000")
MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY", "minioadmin")
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY", "minioadmin")
MINIO_BUCKET = os.getenv("MINIO_BUCKET", "rag-documents")
MINIO_SECURE = os.getenv("MINIO_SECURE", "false").lower() == "true"

# MongoDB setup
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/")
mongo_client = MongoClient(MONGO_URI)
db = mongo_client["chat_app"]
history_collection = db["chat_history"]
users_collection = db["users"]
projects_collection = db["projects"]
requirements_collection = db["requirements"]
versions_collection = db["versions"]
collaborators_collection = db["collaborators"]
api_keys_collection = db["api_keys"]
user_settings_collection = db["user_settings"]
user_settings_collection.create_index([("user", 1)])
user_settings_collection.create_index([("user", 1), ("project_id", 1)])

# Initialize collections for admin and manager modules
admin.users_collection = users_collection
admin.projects_collection = projects_collection
admin.collaborators_collection = collaborators_collection
admin.api_keys_collection = api_keys_collection

manager.users_collection = users_collection
manager.projects_collection = projects_collection
manager.collaborators_collection = collaborators_collection
manager.requirements_collection = requirements_collection
manager.api_keys_collection = api_keys_collection
project_settings_collection = db["project_settings"]

init_admin_collections(
    users_collection, 
    projects_collection, 
    collaborators_collection, 
    api_keys_collection if 'api_keys_collection' in globals() else None
)

# Global variable for MinIO RAG System ONLY
minio_rag_system = None

def initialize_minio_rag():
    """Initialize the MinIO RAG System - REQUIRED FOR APPLICATION TO START"""
    global minio_rag_system
    
    if not MINIO_RAG_AVAILABLE:
        print("❌ CRITICAL ERROR: MinIO RAG System not available")
        print("❌ Application cannot start without MinIO RAG System")
        return False
    
    try:
        PERSIST_DIR = os.getenv("MINIO_PERSIST_DIR", "./chroma_db_minio")
        
        print(f"🤖 Initializing MinIO RAG System...")
        print(f"🗄️ MinIO Endpoint: {MINIO_ENDPOINT}")
        print(f"📦 Bucket: {MINIO_BUCKET}")
        print(f"🧠 OpenRouter Model: {os.getenv('OPENROUTER_MODEL', 'anthropic/claude-3.5-haiku')}")
        
        minio_rag_system = MinIORAGSystem(
            minio_endpoint=MINIO_ENDPOINT,
            minio_access_key=MINIO_ACCESS_KEY,
            minio_secret_key=MINIO_SECRET_KEY,
            bucket_name=MINIO_BUCKET,
            persist_directory=PERSIST_DIR,
            secure=MINIO_SECURE
        )
        
        # Check if documents need to be processed
        stats = minio_rag_system.get_database_stats()
        if stats.get('total_chunks', 0) == 0:
            print("📝 Processing documents from MinIO...")
            if minio_rag_system.process_documents_from_minio():
                print("✅ MinIO RAG System initialized successfully!")
            else:
                print("⚠️ No documents found in MinIO bucket - ready for uploads")
        else:
            print("✅ MinIO RAG System initialized with existing database!")
        
        return True
        
    except Exception as e:
        print(f"❌ CRITICAL ERROR: Failed to initialize MinIO RAG: {e}")
        print("❌ Application cannot start without MinIO RAG System")
        minio_rag_system = None
        return False

def check_rag_availability():
    """Check if MinIO RAG system is available and ready"""
    return MINIO_RAG_AVAILABLE and minio_rag_system is not None

# Create indexes
history_collection.create_index([("user", 1)])
history_collection.create_index([("timestamp", -1)])
projects_collection.create_index([("user", 1)])
requirements_collection.create_index([("project_id", 1)])
requirements_collection.create_index([("user", 1)])
versions_collection.create_index([("requirement_id", 1)])
versions_collection.create_index([("timestamp", -1)])
collaborators_collection.create_index([("project_id", 1)])
collaborators_collection.create_index([("email", 1)])
api_keys_collection.create_index([("user", 1)])

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if request.method == "OPTIONS":
            return "", 200
        
        if "user" not in session:
            return jsonify({"error": "Unauthorized"}), 401
        return f(*args, **kwargs)
    return decorated_function

def is_admin(username):
    """Check if a user has admin role"""
    user = users_collection.find_one({"username": username})
    return user and user.get("role") == "admin"

def is_manager_or_admin(username):
    """Check if a user has manager or admin role"""
    user = users_collection.find_one({"username": username})
    return user and user.get("role") in ["manager", "admin"]

def can_create_projects(username):
    """Check if a user can create projects (manager or admin only)"""
    return is_manager_or_admin(username)

def manager_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user" not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        if not is_manager_or_admin(session["user"]):
            return jsonify({"error": "Manager or admin access required"}), 403
            
        return f(*args, **kwargs)
    return decorated_function

def get_user_api_key(username, project_id=None):
    """Get a user's API key, with optional project-specific override."""
    if project_id:
        project_key = api_keys_collection.find_one({
            "user": username,
            "project_id": project_id
        })
        if project_key and project_key.get("api_key"):
            return project_key["api_key"]
    
    user_key = api_keys_collection.find_one({
        "user": username,
        "project_id": {"$exists": False}
    })
    if user_key and user_key.get("api_key"):
        return user_key["api_key"]
    
    return None

def get_anthropic_client(username, project_id=None):
    """Get an Anthropic client using the appropriate API key."""
    api_key = get_user_api_key(username, project_id)
    if not api_key:
        raise ValueError("No API key available")
    return anthropic.Anthropic(api_key=api_key)


def detect_priority(description):
    if not description:
        return "medium"
    
    description_lower = description.lower()
    
    high_keywords = [
        "security", "authentication", "authorization", "critical", "urgent", "payment", 
        "billing", "data protection", "privacy", "compliance", "regulation", "gdpr",
        "safety", "emergency", "backup", "recovery", "must", "required", "essential",
        "login", "access control", "encryption", "vulnerability", "exploit"
    ]
    
    low_keywords = [
        "nice to have", "optional", "cosmetic", "ui enhancement", "color", "font",
        "styling", "aesthetic", "minor improvement", "suggestion", "preference",
        "nice-to-have", "future enhancement", "wishlist", "optional feature"
    ]
    
    high_count = sum(1 for keyword in high_keywords if keyword in description_lower)
    low_count = sum(1 for keyword in low_keywords if keyword in description_lower)
    
    if high_count > low_count and high_count >= 1:
        return "high"
    elif low_count > high_count and low_count >= 1:
        return "low"
    else:
        return "medium"

def get_priority_label(priority):
    priority_labels = {
        "low": "Faible",
        "medium": "Moyenne", 
        "high": "Élevée"
    }
    return priority_labels.get(priority, priority)

def get_category_label(category):
    category_labels = {
        "functionality": "Fonctionnalité",
        "performance": "Performance",
        "security": "Sécurité",
        "usability": "Utilisabilité",
        "compatibility": "Compatibilité"
    }
    return category_labels.get(category, category)

def get_status_label(status):
    status_labels = {
        "draft": "Brouillon",
        "in_review": "En révision",
        "approved": "Approuvé",
        "implemented": "Implémenté",
        "rejected": "Rejeté"
    }
    return status_labels.get(status, status)

limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["300 per minute"] 
)

# MinIO Configuration - REQUIRED
MINIO_ENDPOINT = os.getenv("MINIO_ENDPOINT", "localhost:9000")
MINIO_ACCESS_KEY = os.getenv("MINIO_ACCESS_KEY", "minioadmin")
MINIO_SECRET_KEY = os.getenv("MINIO_SECRET_KEY", "minioadmin")
MINIO_BUCKET = os.getenv("MINIO_BUCKET", "rag-documents")
MINIO_SECURE = os.getenv("MINIO_SECURE", "false").lower() == "true"

# MongoDB setup
MONGO_URI = os.getenv("MONGO_URI", "mongodb://localhost:27017/")
mongo_client = MongoClient(MONGO_URI)
db = mongo_client["chat_app"]
history_collection = db["chat_history"]
users_collection = db["users"]
projects_collection = db["projects"]
requirements_collection = db["requirements"]
versions_collection = db["versions"]
collaborators_collection = db["collaborators"]
api_keys_collection = db["api_keys"]
user_settings_collection = db["user_settings"]
user_settings_collection.create_index([("user", 1)])
user_settings_collection.create_index([("user", 1), ("project_id", 1)])

# Initialize collections for admin and manager modules
admin.users_collection = users_collection
admin.projects_collection = projects_collection
admin.collaborators_collection = collaborators_collection
admin.api_keys_collection = api_keys_collection

manager.users_collection = users_collection
manager.projects_collection = projects_collection
manager.collaborators_collection = collaborators_collection
manager.requirements_collection = requirements_collection
manager.api_keys_collection = api_keys_collection
project_settings_collection = db["project_settings"]

init_admin_collections(
    users_collection, 
    projects_collection, 
    collaborators_collection, 
    api_keys_collection if 'api_keys_collection' in globals() else None
)

# Global variable for MinIO RAG System ONLY
minio_rag_system = None

def initialize_minio_rag():
    """Initialize the MinIO RAG System - REQUIRED FOR APPLICATION TO START"""
    global minio_rag_system
    
    if not MINIO_RAG_AVAILABLE:
        print("❌ CRITICAL ERROR: MinIO RAG System not available")
        print("❌ Application cannot start without MinIO RAG System")
        return False
    
    try:
        OLLAMA_MODEL = os.getenv("RAG_OLLAMA_MODEL", "qwen3:8b")
        OLLAMA_BASE_URL = os.getenv("RAG_OLLAMA_BASE_URL", "http://35.173.131.200:11434")
        PERSIST_DIR = os.getenv("MINIO_PERSIST_DIR", "./chroma_db_minio")
        
        print(f"🤖 Initializing MinIO RAG System...")
        print(f"🗄️ MinIO Endpoint: {MINIO_ENDPOINT}")
        print(f"📦 Bucket: {MINIO_BUCKET}")
        print(f"🔗 Ollama URL: {OLLAMA_BASE_URL}")
        print(f"🧠 Model: {OLLAMA_MODEL}")
        print(f"💾 Persist Directory: {PERSIST_DIR}")
        
        minio_rag_system = MinIORAGSystem(
            minio_endpoint=MINIO_ENDPOINT,
            minio_access_key=MINIO_ACCESS_KEY,
            minio_secret_key=MINIO_SECRET_KEY,
            bucket_name=MINIO_BUCKET,
            persist_directory=PERSIST_DIR,
            ollama_base_url=OLLAMA_BASE_URL,
            ollama_model=OLLAMA_MODEL,
            secure=MINIO_SECURE
        )
        
        # Check if documents need to be processed
        stats = minio_rag_system.get_database_stats()
        if stats.get('total_chunks', 0) == 0:
            print("📝 Processing documents from MinIO...")
            if minio_rag_system.process_documents_from_minio():
                print("✅ MinIO RAG System initialized successfully!")
            else:
                print("⚠️ No documents found in MinIO bucket - ready for uploads")
        else:
            print("✅ MinIO RAG System initialized with existing database!")
        
        return True
        
    except Exception as e:
        print(f"❌ CRITICAL ERROR: Failed to initialize MinIO RAG: {e}")
        print("❌ Application cannot start without MinIO RAG System")
        minio_rag_system = None
        return False

def check_rag_availability():
    """Check if MinIO RAG system is available and ready"""
    return MINIO_RAG_AVAILABLE and minio_rag_system is not None

# Create indexes
history_collection.create_index([("user", 1)])
history_collection.create_index([("timestamp", -1)])
projects_collection.create_index([("user", 1)])
requirements_collection.create_index([("project_id", 1)])
requirements_collection.create_index([("user", 1)])
versions_collection.create_index([("requirement_id", 1)])
versions_collection.create_index([("timestamp", -1)])
collaborators_collection.create_index([("project_id", 1)])
collaborators_collection.create_index([("email", 1)])
api_keys_collection.create_index([("user", 1)])

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if request.method == "OPTIONS":
            return "", 200
        
        if "user" not in session:
            return jsonify({"error": "Unauthorized"}), 401
        return f(*args, **kwargs)
    return decorated_function

def is_admin(username):
    """Check if a user has admin role"""
    user = users_collection.find_one({"username": username})
    return user and user.get("role") == "admin"

def is_manager_or_admin(username):
    """Check if a user has manager or admin role"""
    user = users_collection.find_one({"username": username})
    return user and user.get("role") in ["manager", "admin"]

def can_create_projects(username):
    """Check if a user can create projects (manager or admin only)"""
    return is_manager_or_admin(username)

def manager_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user" not in session:
            return jsonify({"error": "Unauthorized"}), 401
        
        if not is_manager_or_admin(session["user"]):
            return jsonify({"error": "Manager or admin access required"}), 403
            
        return f(*args, **kwargs)
    return decorated_function

def get_user_api_key(username, project_id=None):
    """Get a user's API key, with optional project-specific override."""
    if project_id:
        project_key = api_keys_collection.find_one({
            "user": username,
            "project_id": project_id
        })
        if project_key and project_key.get("api_key"):
            return project_key["api_key"]
    
    user_key = api_keys_collection.find_one({
        "user": username,
        "project_id": {"$exists": False}
    })
    if user_key and user_key.get("api_key"):
        return user_key["api_key"]
    
    return None

def get_anthropic_client(username, project_id=None):
    """Get an Anthropic client using the appropriate API key."""
    api_key = get_user_api_key(username, project_id)
    if not api_key:
        raise ValueError("No API key available")
    return anthropic.Anthropic(api_key=api_key)

def generate_test_cases_with_minio_rag(requirements, context=""):
    """Generate test cases using the MinIO RAG system"""
    global minio_rag_system
    
    if not minio_rag_system:
        raise ValueError("MinIO RAG system not available")
    
    try:
        full_requirement = f"{requirements}"
        if context:
            full_requirement = f"Context: {context}\n\nRequirement: {requirements}"
        
        print(f"🤖 Generating test cases with MinIO RAG...")
        test_cases = minio_rag_system.generate_test_cases(full_requirement, context)
        
        return test_cases
        
    except Exception as e:
        print(f"Error generating test cases with MinIO RAG: {e}")
        raise

def detect_priority(description):
    if not description:
        return "medium"
    
    description_lower = description.lower()
    
    high_keywords = [
        "security", "authentication", "authorization", "critical", "urgent", "payment", 
        "billing", "data protection", "privacy", "compliance", "regulation", "gdpr",
        "safety", "emergency", "backup", "recovery", "must", "required", "essential",
        "login", "access control", "encryption", "vulnerability", "exploit"
    ]
    
    low_keywords = [
        "nice to have", "optional", "cosmetic", "ui enhancement", "color", "font",
        "styling", "aesthetic", "minor improvement", "suggestion", "preference",
        "nice-to-have", "future enhancement", "wishlist", "optional feature"
    ]
    
    high_count = sum(1 for keyword in high_keywords if keyword in description_lower)
    low_count = sum(1 for keyword in low_keywords if keyword in description_lower)
    
    if high_count > low_count and high_count >= 1:
        return "high"
    elif low_count > high_count and low_count >= 1:
        return "low"
    else:
        return "medium"

def get_priority_label(priority):
    priority_labels = {
        "low": "Faible",
        "medium": "Moyenne", 
        "high": "Élevée"
    }
    return priority_labels.get(priority, priority)

def get_category_label(category):
    category_labels = {
        "functionality": "Fonctionnalité",
        "performance": "Performance",
        "security": "Sécurité",
        "usability": "Utilisabilité",
        "compatibility": "Compatibilité"
    }
    return category_labels.get(category, category)

def get_status_label(status):
    status_labels = {
        "draft": "Brouillon",
        "in_review": "En révision",
        "approved": "Approuvé",
        "implemented": "Implémenté",
        "rejected": "Rejeté"
    }
    return status_labels.get(status, status)
@app.route("/login", methods=["POST"])
def login():
    data = request.json
    username = data.get("username")
    password = data.get("password")

    user = users_collection.find_one({"username": username, "password": password})
    if user:
        session["user"] = username
        session.permanent = True
        return jsonify({
            "message": "Login successful", 
            "username": username,
            "email": user.get("email") or username,
            "role": user.get("role", "user"),
            "is_admin": user.get("role") == "admin",
            "is_manager": user.get("role") in ["manager", "admin"],
            "can_create_projects": user.get("role") in ["manager", "admin"]
        })
    return jsonify({"error": "Invalid credentials"}), 401

@app.route("/logout", methods=["POST"])
@login_required
def logout():
    session.clear()
    return jsonify({"message": "Logged out successfully"})

@app.route("/check_session", methods=["GET", "OPTIONS"])
def check_session():
    if request.method == "OPTIONS":
        return jsonify({}), 200

    if "user" in session:
        user = users_collection.find_one({"username": session["user"]})
        if user:
            user_role = user.get("role", "user")
            return jsonify({
                "authenticated": True,  
                "logged_in": True,      
                "username": session["user"],
                "email": user.get("email") or session["user"],
                "role": user_role,
                "is_admin": user_role == "admin",
                "is_manager": user_role in ["manager", "admin"],
                "can_create_projects": user_role in ["manager", "admin"]
            }), 200
        else:
            session.clear()
            return jsonify({"authenticated": False, "logged_in": False}), 401
    
    return jsonify({"authenticated": False, "logged_in": False}), 401


@app.route("/generate_test_cases", methods=["POST", "OPTIONS"])
@login_required
def generate_test_cases():
    """Generate test cases using project-level LLM settings ONLY"""
    try:
        data = request.json
        requirements = data.get("requirements", "")
        format_type = data.get("format_type", "default")
        context = data.get("context", "")
        example_case = data.get("example_case", "")
        project_id = data.get("project_id", "")
        requirement_id = data.get("requirement_id", "")
        requirement_title = data.get("requirement_title", "")
        
        if not requirements:
            return jsonify({"error": "No requirements provided"}), 400
        
        username = session["user"]
        
        print(f"Generating test cases for user: {username}")
        print(f"Project ID: {project_id}")
        
        # Determine which AI service to use based on PROJECT settings ONLY
        effective_service = get_effective_llm_for_project(project_id) if project_id else None
        
        if not effective_service:
            if not project_id:
                return jsonify({
                    "error": "No project specified.",
                    "suggestion": "Please select a project to generate test cases.",
                    "details": {
                        "project_configured": False,
                        "action_required": "Select project"
                    }
                }), 400
            else:
                return jsonify({
                    "error": "No AI service configured for this project.",
                    "suggestion": "Please ask your project manager to configure the LLM settings for this project.",
                    "details": {
                        "project_configured": False,
                        "action_required": "Contact project manager",
                        "message": "Only project managers can configure AI settings."
                    }
                }), 400
        
        print(f"Using AI Service: {effective_service}")
        print(f"Requirements: {requirements[:50]}...")
        
        # Generate test cases based on effective service
        if effective_service == "claude":
            api_key = get_project_api_key(project_id)
            if not api_key:
                return jsonify({
                    "error": "Claude API key not configured for this project.",
                    "suggestion": "Please ask your project manager to add the Claude API key.",
                    "details": {
                        "action_required": "Contact project manager"
                    }
                }), 400
            
            # Generate test cases using Claude with project API key
            test_cases = generate_test_cases_claude(
                requirements, format_type, context, example_case, 
                requirement_id, requirement_title, api_key, project_id
            )
        else:
            # Generate test cases using MinIO RAG
            test_cases = generate_test_cases_minio_rag(
                requirements, format_type, context, example_case,
                requirement_id, requirement_title, project_id
            )
        
        return jsonify({
            "test_cases": test_cases,
            "service_used": effective_service,
            "project_configured": True,
            "configured_by": "project_manager"
        })
        
    except Exception as e:
        print(f"Error generating test cases: {e}")
        return jsonify({"error": str(e)}), 500


def generate_test_cases_claude(requirements, format_type, context, example_case, requirement_id, requirement_title, api_key, project_id=None):
    """Generate test cases using Claude API with project-level API key and language"""
    try:
        # Get project language if project_id is provided
        project_language = "en"  # default
        if project_id:
            project = projects_collection.find_one({"id": project_id})
            if project:
                project_language = project.get('language', 'en')
        
        # Create Anthropic client with project API key
        anthropic_client = anthropic.Anthropic(api_key=api_key)
        
        test_case_instruction = generate_test_case_prompt(requirements, format_type, context, example_case, project_language)
       
        # Call Claude API
        response = anthropic_client.messages.create(
            model="claude-3-5-haiku-20241022",
            max_tokens=2000,
            messages=[{"role": "user", "content": test_case_instruction}]
        )
        
        test_cases = response.content[0].text
        
        # Save to history with project context
        username = session.get("user")
        
        # ✅ FIXED: Mark as "generated" not "ai_assistant"
        history_entry = {
            "user": username,
            "project_id": project_id,
            "requirement_id": requirement_id,
            "requirement_title": requirement_title,
            "requirements": requirements,
            "test_cases": test_cases,
            "format_type": format_type,
            "context": context,
            "example_case": example_case,
            "ai_service": "claude",  
            "language": project_language,
            "timestamp": datetime.now(timezone.utc),
            "used_project_settings": True,
            "update_type": "generated",  # ✅ Keep as "generated"
            "source": "initial_generation",  # ✅ Mark as initial generation
            "generation_method": "api_call"  # ✅ Distinguish from chat
        }
        
        if history_collection is not None:
            history_collection.insert_one(history_entry)
        
        return test_cases
        
    except Exception as e:
        print(f"Claude API error: {e}")
        raise Exception(f"Claude API error: {str(e)}")

def generate_test_cases_minio_rag(requirements, format_type, context, example_case, requirement_id, requirement_title, project_id=None):
    """Generate test cases using MinIO RAG system with language support"""
    try:
        if not check_rag_availability():
            raise Exception("MinIO RAG system not available")
        
        # Get project language if project_id is provided
        project_language = "en"  # default
        if project_id:
            project = projects_collection.find_one({"id": project_id})
            if project:
                project_language = project.get('language', 'en')
        
        # Create a formatted prompt that works with your RAG system
        full_context = ""
        if context:
            full_context = f"Context: {context}\n\n"
        
        # Add language instruction
        if project_language == 'fr':
            language_instruction = "Veuillez répondre en français.\n\n"
        else:
            language_instruction = "Please respond in English.\n\n"
        
        full_requirement = f"{language_instruction}{full_context}Requirement: {requirements}"
        
        # Call MinIO RAG system
        print(f"🤖 Calling MinIO RAG generate_test_cases method...")
        test_cases = minio_rag_system.generate_test_cases(full_requirement, context)
        
        # Save to history with project context
        username = session.get("user")
        
        # ✅ FIXED: Mark as "generated" not "ai_assistant"
        history_entry = {
            "user": username,
            "project_id": project_id,
            "requirement_id": requirement_id,
            "requirement_title": requirement_title,
            "requirements": requirements,
            "test_cases": test_cases,
            "format_type": format_type,
            "context": context,
            "example_case": example_case,
            "ai_service": "minio",  
            "language": project_language,
            "timestamp": datetime.now(timezone.utc),
            "used_project_settings": True,
            "update_type": "generated",  # ✅ Keep as "generated"
            "source": "initial_generation",  # ✅ Mark as initial generation
            "generation_method": "api_call"  # ✅ Distinguish from chat
        }
        
        if history_collection is not None:
            history_collection.insert_one(history_entry)
        
        return test_cases
        
    except Exception as e:
        print(f"MinIO RAG error: {e}")
        raise Exception(f"MinIO RAG error: {str(e)}")

@app.route("/generate_test_cases_stream", methods=["POST"])
@login_required
@limiter.limit("5 per minute")
def generate_test_cases_stream():
    data = request.json
    requirements = data.get("requirements", "")
    format_type = data.get("format_type", "default")
    context = data.get("context", "")
    example_case = data.get("example_case", "")
    project_id = data.get("project_id", "")
    
    if not requirements:
        return jsonify({"error": "No requirements provided"}), 400
    
    if not project_id:
        return jsonify({"error": "Project ID is required"}), 400
    
    # Get project language
    project_language = "en"  # default
    project = projects_collection.find_one({"id": project_id})
    if project:
        project_language = project.get('language', 'en')
    
    # Generate test case prompt
    test_case_instruction = generate_test_case_prompt(requirements, format_type, context, example_case, project_language)
    username = session["user"]
    
    def generate():
        try:
            full_response = ""
            
            # Use project-configured AI service - CLAUDE OR MINIO ONLY
            effective_service = get_effective_llm_for_project(project_id)
            
            if effective_service == "claude":
                # Use Claude with project API key
                api_key = get_project_api_key(project_id)
                if not api_key:
                    yield f"data: {json.dumps({'error': 'Claude API not configured for this project'})}\n\n"
                    return
                
                client = anthropic.Anthropic(api_key=api_key)
                with client.messages.stream(
                    model="claude-3-5-sonnet-20241022",
                    max_tokens=4000,
                    messages=[{"role": "user", "content": test_case_instruction}]
                ) as stream:
                    for text in stream.text_stream:
                        full_response += text
                        yield f"data: {text}\n\n"
            
            elif effective_service == "minio":
                # Use MinIO RAG System
                if check_rag_availability():
                    response = minio_rag_system.generate_test_cases(test_case_instruction)
                    full_response = response
                    yield f"data: {response}\n\n"
                else:
                    yield f"data: {json.dumps({'error': 'MinIO RAG system not available'})}\n\n"
                    return
            
            else:
                # No valid service configured
                yield f"data: {json.dumps({'error': 'No AI service configured. Please contact your manager to configure Claude API or ensure MinIO RAG is available.'})}\n\n"
                return
            
            # Save to history with project context
            history_entry = {
                "user": username,
                "project_id": project_id,
                "requirements": requirements,
                "test_cases": full_response,
                "format_type": format_type,
                "context": context,
                "example_case": example_case,
                "ai_service": effective_service,
                "language": project_language,
                "timestamp": datetime.now(timezone.utc),
                "used_project_settings": True,
                "update_type": "generated",  # ✅ Mark as generated
                "source": "streaming_generation",  # ✅ Mark as streaming generation
                "generation_method": "stream_api"  # ✅ Distinguish from chat
            }
            
            if history_collection is not None:
                history_collection.insert_one(history_entry)
            
            yield "data: [DONE]\n\n"
            
        except Exception as e:
            print(f"Error in test case generation: {e}")
            yield f"data: {json.dumps({'error': f'Error generating test cases: {str(e)}'})}\n\n"
    
    return Response(generate(), content_type="text/event-stream", headers={
        'Cache-Control': 'no-cache',
        'Connection': 'keep-alive'
    })




def generate_test_case_prompt(requirements, format_type="default", context="", example_case="", language="en"):
    """Generate the prompt for test case generation with language support"""
    
    if language == "fr":
        # French prompts
        if format_type == "custom" and example_case.strip():
            example_format = example_case
        elif format_type == "gherkin":
            example_format = example_case if example_case.strip() else "Format Gherkin"
        else:
            example_format = """Cas de Test 1: Connexion Valide
Description: L'utilisateur se connecte avec des identifiants valides
Prérequis: L'utilisateur a un compte valide
Étapes:
    1. Naviguer vers la page de connexion.
    2. Saisir un email et mot de passe valides.
    3. Cliquer sur "Se connecter".
Résultat Attendu: L'utilisateur est connecté avec succès et redirigé vers le tableau de bord.

Cas de Test 2: Connexion Invalide
Description: L'utilisateur tente de se connecter avec des identifiants invalides
Prérequis: L'utilisateur est sur la page de connexion
Étapes:
    1. Accéder à la page de connexion.
    2. Saisir un email valide et un mot de passe invalide.
    3. Cliquer sur "Se connecter".
Résultat Attendu: Un message d'erreur s'affiche et l'utilisateur reste sur la page de connexion.
"""

        instruction = f"""
Générez des cas de test pour l'exigence suivante en utilisant le format spécifié.
{"Contexte fonctionnel: " + context if context else ""} 
Exigence: {requirements}
Format:
{example_format}

Concentrez-vous sur :
1. Les scénarios de test positifs et négatifs
2. Les cas limites et les conditions aux bornes
3. Des étapes de test claires et réalisables
4. Les résultats attendus pour chaque cas de test
5. Les prérequis et les exigences de données de test

Fournissez au moins 5-8 cas de test complets couvrant différents scénarios.
"""
    else:
        # English prompts (your existing logic)
        example_format_default = """Test Case 1: Valid Login
Description: User logs in with valid credentials
Preconditions: User has a valid account
Steps:
    1. Navigate to the login page.
    2. Enter valid email and password.
    3. Click on "Login".
Expected Result: User is successfully logged in and redirected to the dashboard.

Test Case 2: Invalid Login
Description: User attempts to log in with invalid credentials
Preconditions: User is on the login page
Steps:
    1. Access the login page.
    2. Enter valid email and invalid password.
    3. Click on "Login".
Expected Result: An error message is displayed, and the user remains on the login page.
"""

        if format_type == "custom" and example_case.strip():
            example_format = example_case
        elif format_type == "gherkin":
            example_format = example_case if example_case.strip() else "Gherkin format"
        else:
            example_format = example_format_default

        instruction = f"""
Generate test cases for the following requirement using the specified format.
{"Functional context: " + context if context else ""} 
Requirement: {requirements}
Format:
{example_format}

Focus on:
1. Both positive and negative test scenarios
2. Edge cases and boundary conditions
3. Clear, actionable test steps
4. Expected results for each test case
5. Prerequisites and test data requirements

Provide at least 5-8 comprehensive test cases covering different scenarios.
"""
    
    return instruction

# Additional helper functions that you might need to implement

def track_project_context(project_id):
    """Helper function to track current project context in session"""
    if project_id:
        session["current_project_id"] = project_id

# You might want to add this to your project selection endpoints
@app.route("/select_project/<project_id>", methods=["POST"])
@login_required
def select_project(project_id):
    """Set the current project context for the user session"""
    try:
        username = session["user"]
        
        # Verify user has access to this project
        project = projects_collection.find_one({"id": project_id})
        if not project:
            return jsonify({"error": "Project not found"}), 404
        
        # Check if user is owner or collaborator
        user = users_collection.find_one({"username": username})
        user_role = user.get("role", "user")
        
        if (user_role in ["manager", "admin"] and project["user"] == username) or \
           (username in project.get("collaborators", [])):
            
            session["current_project_id"] = project_id
            
            # Get project LLM settings for frontend
            project_settings = get_project_llm_settings(project_id)
            effective_service = get_effective_llm_for_project(project_id)
            
            return jsonify({
                "message": "Project selected successfully",
                "project": {
                    "id": project_id,
                    "name": project.get("name"),
                    "effective_service": effective_service,
                    "has_llm_settings": bool(project_settings)
                }
            })
        else:
            return jsonify({"error": "Access denied to this project"}), 403
            
    except Exception as e:
        print(f"Error selecting project: {e}")
        return jsonify({"error": str(e)}), 500

# Update the main test case generation to use project context from request
# instead of session if project_id is provided in the request data
@app.route("/check_api_key_status", methods=["GET"])
@login_required
def check_api_key_status():
    """Check if project has API keys configured - MANAGER ACCESS ONLY"""
    username = session["user"]
    project_id = request.args.get("project_id")
    
    try:
        # Check if user is manager/admin
        user = users_collection.find_one({"username": username})
        if not user or user.get("role") not in ["manager", "admin"]:
            return jsonify({
                "error": "Only managers can view API key status",
                "user_role": user.get("role", "user") if user else "user"
            }), 403
        
        if not project_id:
            return jsonify({
                "has_project_api_key": False,
                "message": "No project selected"
            })
        
        # Check project-level API key only
        project_settings = get_project_llm_settings(project_id)
        has_project_key = bool(project_settings and project_settings.get("claude_api_key"))
        
        status = {
            "has_project_api_key": has_project_key,
            "project_configured": bool(project_settings),
            "effective_service": get_effective_llm_for_project(project_id)
        }
        
        return jsonify(status)
        
    except Exception as e:
        print(f"Error checking API key status: {e}")
        return jsonify({"error": str(e)}), 500
@app.route("/save_test_cases", methods=["POST"])
@login_required
def save_test_cases():
    data = request.json
    test_cases = data.get("test_cases", "")
    requirements = data.get("requirements", "")
    project_id = data.get("project_id", "")
    requirement_id = data.get("requirement_id", "")
    requirement_title = data.get("requirement_title", "")
    
    username = session["user"]
    current_time = datetime.now(timezone.utc)
    
    history_data = {
        "user": username,
        "test_cases": test_cases,
        "timestamp": current_time,
        "requirements": requirements,
        "context": "",
        "project_id": project_id,
        "update_type": "manual_edit"
    }
    
    if requirement_id:
        history_data["requirement_id"] = requirement_id
    if requirement_title:
        history_data["requirement_title"] = requirement_title
    
    history_collection.insert_one(history_data)
    
    return jsonify({
        "message": "Test cases saved successfully",
        "timestamp": current_time.isoformat()
    })


@app.route("/history", methods=["GET"])
@login_required
def get_history():
    username = session["user"]
    project_id = request.args.get("project_id")
    
    query = {"user": username}
    if project_id:
        query["project_id"] = project_id
    
    history = list(history_collection.find(query).sort("timestamp", -1).limit(50))
    
    for item in history:
        item["_id"] = str(item["_id"])
    
    return jsonify({"history": history})

@app.route("/history/<history_id>", methods=["GET"])
@login_required
def get_history_item(history_id):
    username = session["user"]
    
    try:
        object_id = ObjectId(history_id)
    except:
        return jsonify({"error": "Invalid history ID"}), 400
    
    item = history_collection.find_one({
        "_id": object_id,
        "user": username
    })
    
    if not item:
        return jsonify({"error": "History item not found"}), 404
    
    item["_id"] = str(item["_id"])
    return jsonify({"item": item})

@app.route("/update_test_cases/<history_id>", methods=["PUT"])
@login_required
def update_test_cases(history_id):
    data = request.json
    test_cases = data.get("test_cases", "")
    requirements = data.get("requirements", "")
    project_id = data.get("project_id", "")
    requirement_id = data.get("requirement_id", "")
    requirement_title = data.get("requirement_title", "")
    update_type = data.get("update_type", "manual_edit")
    
    username = session["user"]
    current_time = datetime.now(timezone.utc)
    
    try:
        object_id = ObjectId(history_id)
    except:
        return jsonify({"error": "Invalid history ID"}), 400
    
    # Find the existing history item
    existing_item = history_collection.find_one({
        "_id": object_id,
        "user": username
    })
    
    if not existing_item:
        return jsonify({"error": "History item not found or access denied"}), 404
    
    # Update the existing history item
    update_data = {
        "test_cases": test_cases,
        "timestamp": current_time,
        "update_type": update_type
    }
    
    if requirements:
        update_data["requirements"] = requirements
    if project_id:
        update_data["project_id"] = project_id
    if requirement_id:
        update_data["requirement_id"] = requirement_id
    if requirement_title:
        update_data["requirement_title"] = requirement_title
    
    history_collection.update_one(
        {"_id": object_id},
        {"$set": update_data}
    )
    
    return jsonify({
        "message": "Test cases updated successfully",
        "timestamp": current_time.isoformat()
    })

@app.route("/history/<history_id>", methods=["DELETE"])
@login_required
def delete_history_item(history_id):
    username = session["user"]
    
    try:
        object_id = ObjectId(history_id)
    except:
        return jsonify({"error": "Invalid history ID"}), 400
    
    result = history_collection.delete_one({
        "_id": object_id,
        "user": username
    })
    
    if result.deleted_count == 0:
        return jsonify({"error": "History item not found"}), 404
    
    return jsonify({"message": "History item deleted successfully"})


@app.route("/chat_with_assistant", methods=["POST"])
@login_required
def chat_with_test_cases():
    data = request.json
    user_message = data.get("message", "")
    test_cases = data.get("test_cases", "")
    project_id = data.get("project_id", "")
    requirement_id = data.get("requirement_id", "")
    chat_history = data.get("chat_history", [])
    direct_mode = data.get("direct_mode", False)
    active_history_id = data.get("active_history_id")
    
    username = session["user"]
    
    # Enhanced logging
    print(f"Chat request received from user: {username}")
    print(f"Message: {user_message[:100]}..." if len(user_message) > 100 else f"Message: {user_message}")
    print(f"Direct mode: {direct_mode}, Active history ID: {active_history_id}")
    
    # REQUIRE PROJECT ID
    if not project_id:
        error_msg = "Project ID is required for AI chat assistant"
        print(f"ERROR: {error_msg}")
        return Response(
            f"data: {json.dumps({'error': error_msg})}\n\n", 
            content_type="text/event-stream",
            headers={'Cache-Control': 'no-cache', 'Connection': 'keep-alive'}
        )
    
    # Get the effective LLM service for this project (Claude or MinIO only)
    effective_service = get_effective_llm_for_project(project_id)
    print(f"Using {effective_service} service for project {project_id}")
    
    if not effective_service:
        error_msg = "No AI service configured for this project. Please contact your manager to configure Claude API or ensure MinIO RAG is available."
        print(f"ERROR: {error_msg}")
        return Response(
            f"data: {json.dumps({'error': error_msg})}\n\n", 
            content_type="text/event-stream",
            headers={'Cache-Control': 'no-cache', 'Connection': 'keep-alive'}
        )
    
    # Create context for the AI based on mode
    if direct_mode:
        context_parts = [
            "You are a test case assistant. Your primary job is to directly modify test cases based on user requests.",
            "CRITICAL INSTRUCTIONS:",
            "1. When the user asks for changes, you MUST output the COMPLETE updated test cases in a code block.",
            "2. Use the exact format: ```testcases followed by the complete test cases, then ```",
            "3. Include ALL test cases in your output, not just the modified ones.",
            "4. After the code block, add EXACTLY this message: 'Modifications appliquées avec succès.'",
            "5. DO NOT explain what changes you're making beforehand - show the complete updated test cases immediately."
        ]
    else:
        context_parts = [
            "You are a helpful test case assistant. You can help users understand, analyze, and modify test cases.",
            "You should be conversational and helpful while maintaining accuracy."
        ]
    
    if test_cases:
        context_parts.append(f"\nCurrent test cases:\n{test_cases}")
    
    context = "\n".join(context_parts)
    
    # Prepare messages for the AI
    messages = []
    
    # Add system context
    messages.append({"role": "system", "content": context})
    
    # Add chat history (limit to last 10 messages to avoid token limits)
    recent_history = chat_history[-10:] if len(chat_history) > 10 else chat_history
    for msg in recent_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
    
    # Add current user message
    messages.append({"role": "user", "content": user_message})
    
    def generate():
        try:
            full_response = ""
            updated_test_cases = None
            
            # ONLY CLAUDE OR MINIO RAG
            if effective_service == "claude":
                # Use Claude API
                api_key = get_project_api_key(project_id)
                if not api_key:
                    yield f"data: {json.dumps({'error': 'Claude API key not configured for this project. Please contact your manager to configure it.'})}\n\n"
                    yield "data: [DONE]\n\n"
                    return
                
                print(f"Using Claude API for chat")
                
                try:
                    client = anthropic.Anthropic(api_key=api_key)
                    
                    # Convert messages to Claude format (system message separate)
                    system_message = None
                    claude_messages = []
                    
                    for msg in messages:
                        if msg["role"] == "system":
                            system_message = msg["content"]
                        else:
                            claude_messages.append(msg)
                    
                    # Create Claude streaming request
                    with client.messages.stream(
                        model="claude-3-5-sonnet-20241022",
                        max_tokens=4000,
                        system=system_message,
                        messages=claude_messages
                    ) as stream:
                        for text in stream.text_stream:
                            full_response += text
                            yield f"data: {json.dumps({'message': text})}\n\n"
                            
                except Exception as claude_error:
                    print(f"Claude API error: {claude_error}")
                    yield f"data: {json.dumps({'error': f'Claude API error: {str(claude_error)}'})}\n\n"
                    yield "data: [DONE]\n\n"
                    return
                    
            elif effective_service == "minio":
                # Use MinIO RAG System
                try:
                    print(f"Using MinIO RAG for chat")
                    
                    if not check_rag_availability():
                        yield f"data: {json.dumps({'error': 'MinIO RAG system is not available. Please contact your manager to configure Claude API.'})}\n\n"
                        yield "data: [DONE]\n\n"
                        return
                    
                    # For direct mode test case modifications, use generate_test_cases
                    if direct_mode and test_cases and any(keyword in user_message.lower() 
                        for keyword in ['modify', 'change', 'update', 'add', 'remove', 'edit']):
                        
                        # Generate modified test cases using MinIO RAG
                        context_prompt = f"""
                        Current test cases:
                        {test_cases}
                        
                        User request: {user_message}
                        
                        Please modify the test cases according to the user's request and return the COMPLETE updated test cases.
                        """
                        
                        response = minio_rag_system.generate_test_cases(context_prompt)
                        full_response = response
                        updated_test_cases = response
                        
                        yield f"data: {json.dumps({'message': response})}\n\n"
                    else:
                        # Regular chat using MinIO RAG
                        response = minio_rag_system.query(user_message)
                        full_response = response
                        
                        yield f"data: {json.dumps({'message': response})}\n\n"
                        
                except Exception as rag_error:
                    print(f"MinIO RAG error: {rag_error}")
                    yield f"data: {json.dumps({'error': f'MinIO RAG error: {str(rag_error)}'})}\n\n"
                    yield "data: [DONE]\n\n"
                    return
            
            else:
                # No valid service configured
                error_msg = "No valid AI service configured. Please contact your manager to configure Claude API or ensure MinIO RAG is available."
                yield f"data: {json.dumps({'error': error_msg})}\n\n"
                yield "data: [DONE]\n\n"
                return
            
            # POST-PROCESSING FOR TEST CASE UPDATES
            if direct_mode and full_response:
                # Extract test cases from response if they're in a code block
                import re
                test_case_pattern = r'```(?:testcases?)?\s*(.*?)\s*```'
                matches = re.findall(test_case_pattern, full_response, re.DOTALL | re.IGNORECASE)
                
                if matches:
                    # Use the last match as the updated test cases
                    updated_test_cases = matches[-1].strip()
                    print(f"Extracted updated test cases: {len(updated_test_cases)} characters")
                
                # If we found updated test cases, save to history
                if updated_test_cases and active_history_id:
                    try:
                        object_id = ObjectId(active_history_id)
                        current_time = datetime.now(timezone.utc)
                        
                        # ✅ FIXED: Properly mark as AI Assistant modification
                        update_data = {
                            "test_cases": updated_test_cases,
                            "updated_at": current_time,
                            "timestamp": current_time,  # Update main timestamp too
                            "ai_service": effective_service,
                            "update_type": "ai_assistant",     # ✅ Mark as AI Assistant
                            "modified_by_ai": True,           # ✅ Additional flag
                            "last_ai_service": effective_service,
                            "source": "chat_modification",    # ✅ Mark source as chat
                            "generation_method": "chat_interaction"  # ✅ Distinguish from API
                        }
                        
                        history_collection.update_one(
                            {"_id": object_id},
                            {"$set": update_data}
                        )
                        
                        print(f"✅ Updated history {active_history_id} with AI Assistant modification")
                        
                        yield f"data: {json.dumps({
                            'updated_test_cases': updated_test_cases, 
                            'history_updated': True, 
                            'confirmation': 'Modifications appliquées avec succès.',
                            'update_source': 'ai_assistant',  # ✅ Send update source to frontend
                            'service_used': effective_service
                        })}\n\n"
                        
                    except Exception as update_error:
                        print(f"Error updating history: {update_error}")
                        yield f"data: {json.dumps({
                            'updated_test_cases': updated_test_cases, 
                            'history_updated': False, 
                            'update_error': str(update_error), 
                            'confirmation': 'Modifications appliquées avec succès.',
                            'update_source': 'ai_assistant'
                        })}\n\n"
            
        except Exception as e:
            error_msg = f"Unexpected error in AI chat processing: {str(e)}"
            print(f"ERROR: {error_msg}")
            import traceback
            traceback.print_exc()
            yield f"data: {json.dumps({'error': error_msg})}\n\n"
            yield "data: [DONE]\n\n"
    
    return Response(generate(), content_type="text/event-stream", headers={
        'Cache-Control': 'no-cache',
        'Connection': 'keep-alive'
    })


# Project Management
@app.route("/projects", methods=["GET"])
@login_required
def get_projects():
    username = session["user"]
    user = users_collection.find_one({"username": username})
    user_role = user.get("role", "user") if user else "user"
    
    # Build query based on user role
    if user_role in ["manager", "admin"]:
        # Managers and admins can see projects they own or collaborate on
        query = {
            "$or": [
                {"user": username},
                {"collaborators": username}
            ]
        }
    else:
        # Regular users can only see projects where they are collaborators
        query = {"collaborators": username}
    
    projects = list(projects_collection.find(query))
    
    for project in projects:
        project["_id"] = str(project["_id"])
        project["is_owner"] = project["user"] == username
        project["user_role"] = user_role
    
    return jsonify({"projects": projects})

@app.route("/projects", methods=["POST"])
@manager_required
def create_project():
    data = request.json
    username = session["user"]
    
    project = {
        "id": str(uuid.uuid4()),
        "user": username,
        "name": data.get("name"),
        "context": data.get("context", ""),
        "language": data.get("language", "en"),  # Add this line
        "collaborators": [],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by_role": "manager"
    }
    
    # Insert the project and get the _id
    result = projects_collection.insert_one(project)
    
    # Create a copy of the project to return
    response_project = project.copy()
    response_project["_id"] = str(result.inserted_id)
    
    return jsonify({"message": "Project created", "project": response_project})

@app.route("/projects/<project_id>", methods=["GET"])
@login_required
def get_project(project_id):
    username = session["user"]
    user = users_collection.find_one({"username": username})
    user_role = user.get("role", "user") if user else "user"
    
    # Check access based on role
    if user_role in ["manager", "admin"]:
        project = projects_collection.find_one({
            "id": project_id,
            "$or": [
                {"user": username},
                {"collaborators": username}
            ]
        })
    else:
        project = projects_collection.find_one({
            "id": project_id,
            "collaborators": username
        })
    
    if not project:
        return jsonify({"error": "Project not found or access denied"}), 404
    
    project["_id"] = str(project["_id"])
    project["is_owner"] = project["user"] == username
    project["user_role"] = user_role
    
    return jsonify({"project": project})

@app.route("/projects/<project_id>", methods=["PUT"])
@login_required
def update_project(project_id):
    username = session["user"]
    data = request.json
    
    # Get user details to check role
    user = users_collection.find_one({"username": username})
    user_role = user.get("role", "user") if user else "user"
    
    # Find the project
    project = projects_collection.find_one({"id": project_id})
    if not project:
        return jsonify({"error": "Project not found"}), 404
    
    # Check permissions: only owner or admin can edit
    if project["user"] != username and user_role != "admin":
        return jsonify({"error": "You don't have permission to edit this project"}), 403
    
    # Prepare update data
    update_data = {}
    
    # Update allowed fields
    if "name" in data and data["name"].strip():
        update_data["name"] = data["name"].strip()
    
    if "context" in data:
        update_data["context"] = data["context"]
    
    if "language" in data and data["language"] in ["en", "fr"]:
        update_data["language"] = data["language"]
    
    # Add timestamp
    if update_data:
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
        
        # Update the project
        projects_collection.update_one(
            {"id": project_id},
            {"$set": update_data}
        )
        
        print(f"Project {project_id} updated by {username}: {update_data}")
    
    return jsonify({"message": "Project updated successfully"})

@app.route("/projects/<project_id>", methods=["DELETE"])
@login_required
def delete_project(project_id):
    username = session["user"]
    
    project = projects_collection.find_one({
        "id": project_id,
        "user": username
    })
    
    if not project:
        return jsonify({"error": "Project not found or you don't have permission"}), 404
    
    projects_collection.delete_one({"id": project_id})
    requirements_collection.delete_many({"project_id": project_id})
    collaborators_collection.delete_many({"project_id": project_id})
    
    return jsonify({"message": "Project deleted successfully"})

# Requirement Management
@app.route("/projects/<project_id>/requirements", methods=["GET"])
@login_required
def get_requirements(project_id):
    username = session["user"]
    user = users_collection.find_one({"username": username})
    user_role = user.get("role", "user") if user else "user"
    
    # Check access based on role
    if user_role in ["manager", "admin"]:
        # Managers and admins can access projects they own or collaborate on
        project = projects_collection.find_one({
            "id": project_id,
            "$or": [
                {"user": username},
                {"collaborators": username}
            ]
        })
    else:
        # Regular users can only access projects where they are collaborators
        project = projects_collection.find_one({
            "id": project_id,
            "collaborators": username
        })
    
    if not project:
        return jsonify({"error": "Project not found or access denied"}), 404
    
    requirements = list(requirements_collection.find({
        "project_id": project_id
    }))
    
    for req in requirements:
        req["_id"] = str(req["_id"])
    
    return jsonify({"requirements": requirements})

@app.route("/projects/<project_id>/requirements", methods=["POST"])
@login_required
def create_requirement(project_id):
    data = request.json
    username = session["user"]
    
    project = projects_collection.find_one({
        "id": project_id,
        "$or": [
            {"user": username},
            {"collaborators": username}
        ]
    })
    
    if not project:
        return jsonify({"error": "Project not found or access denied"}), 404
    
    # Generate priority automatically based on description content
    description = data.get("description", "")
    auto_priority = detect_priority(description)
    
    # Use the auto-generated priority if none was specified
    priority = data.get("priority") or auto_priority
    
    requirement = {
        "id": str(uuid.uuid4()),
        "user": username,
        "project_id": project_id,
        "title": data.get("title"),
        "description": description,
        "category": data.get("category", "functionality"),
        "priority": priority,
        "status": data.get("status", "draft"),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "priority_auto_generated": priority == auto_priority
    }
    
    result = requirements_collection.insert_one(requirement)
    requirement["_id"] = str(result.inserted_id)
    
    return jsonify({
        "message": "Requirement created", 
        "requirement": requirement,
        "auto_priority_detected": auto_priority
    })

@app.route("/requirements/<requirement_id>", methods=["GET"])
@login_required
def get_requirement(requirement_id):
    username = session["user"]
    
    requirement = requirements_collection.find_one({"id": requirement_id})
    if not requirement:
        return jsonify({"error": "Requirement not found"}), 404
    
    project = projects_collection.find_one({
        "id": requirement["project_id"],
        "$or": [
            {"user": username},
            {"collaborators": username}
        ]
    })
    
    if not project:
        return jsonify({"error": "Access denied"}), 403
    
    requirement["_id"] = str(requirement["_id"])
    return jsonify({"requirement": requirement})

@app.route("/requirements/<requirement_id>", methods=["PUT"])
@login_required
def update_requirement(requirement_id):
    username = session["user"]
    data = request.json
    
    requirement = requirements_collection.find_one({"id": requirement_id})
    if not requirement:
        return jsonify({"error": "Requirement not found"}), 404
    
    project = projects_collection.find_one({
        "id": requirement["project_id"],
        "$or": [
            {"user": username},
            {"collaborators": username}
        ]
    })
    
    if not project:
        return jsonify({"error": "Access denied"}), 403
    
    update_data = {}
    if "title" in data:
        update_data["title"] = data["title"]
    if "description" in data:
        update_data["description"] = data["description"]
    if "category" in data:
        update_data["category"] = data["category"]
    if "priority" in data:
        update_data["priority"] = data["priority"]
    if "status" in data:
        update_data["status"] = data["status"]
    
    if update_data:
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
        requirements_collection.update_one(
            {"id": requirement_id},
            {"$set": update_data}
        )
    
    return jsonify({"message": "Requirement updated successfully"})

@app.route("/requirements/<requirement_id>", methods=["DELETE"])
@login_required
def delete_requirement(requirement_id):
    username = session["user"]
    
    requirement = requirements_collection.find_one({"id": requirement_id})
    if not requirement:
        return jsonify({"error": "Requirement not found"}), 404
    
    project = projects_collection.find_one({
        "id": requirement["project_id"],
        "$or": [
            {"user": username},
            {"collaborators": username}
        ]
    })
    
    if not project:
        return jsonify({"error": "Access denied"}), 403
    
    requirements_collection.delete_one({"id": requirement_id})
    
    return jsonify({"message": "Requirement deleted successfully"})

# Project Collaboration
@app.route("/projects/<project_id>/collaborators", methods=["GET"])
@login_required
def get_collaborators(project_id):
    username = session["user"]
    user = users_collection.find_one({"username": username})
    user_role = user.get("role", "user") if user else "user"
    
    # Check access - owners, collaborators, and admins can view
    project = projects_collection.find_one({"id": project_id})
    if not project:
        return jsonify({"error": "Project not found"}), 404
    
    # Check permissions
    if (project["user"] != username and 
        username not in project.get("collaborators", []) and 
        user_role != "admin"):
        return jsonify({"error": "Access denied"}), 403
    
    # Get collaborators with enhanced details
    collaborators = list(collaborators_collection.find({"project_id": project_id}))
    
    # Enhance with user details
    for collab in collaborators:
        collab["_id"] = str(collab["_id"])
        
        # Get user details if not already present
        if not collab.get("username"):
            user_details = users_collection.find_one({"email": collab["email"]})
            if user_details:
                collab["username"] = user_details["username"]
    
    return jsonify({"collaborators": collaborators})

# Enhanced add collaborator endpoint with better error handling
@app.route("/projects/<project_id>/collaborators", methods=["POST"])
@login_required
def add_collaborator(project_id):
    username = session["user"]
    data = request.json
    email = data.get("email", "").strip().lower()
    
    if not email:
        return jsonify({"error": "Email is required"}), 400
    
    # Get user details to check role
    user = users_collection.find_one({"username": username})
    user_role = user.get("role", "user") if user else "user"
    
    # Find the project
    project = projects_collection.find_one({"id": project_id})
    if not project:
        return jsonify({"error": "Project not found"}), 404
    
    # Check permissions: only owner or admin can add collaborators
    if project["user"] != username and user_role != "admin":
        return jsonify({"error": "Only project owners and admins can add collaborators"}), 403
    
    # Check if user exists
    collaborator_user = users_collection.find_one({"email": email})
    if not collaborator_user:
        return jsonify({"error": "No user found with this email address"}), 404
    
    collaborator_username = collaborator_user["username"]
    
    # Check if user is the owner
    if email == user.get("email", "").lower():
        return jsonify({"error": "Cannot add project owner as collaborator"}), 400
    
    # Check if already a collaborator
    existing_collab = collaborators_collection.find_one({
        "project_id": project_id,
        "email": email
    })
    
    if existing_collab:
        return jsonify({"error": "User is already a collaborator on this project"}), 400
    
    # Add to collaborators collection
    collaborator = {
        "project_id": project_id,
        "email": email,
        "username": collaborator_username,
        "added_by": username,
        "added_at": datetime.now(timezone.utc).isoformat()
    }
    
    result = collaborators_collection.insert_one(collaborator)
    
    # Update project collaborators list
    projects_collection.update_one(
        {"id": project_id},
        {"$addToSet": {"collaborators": collaborator_username}}
    )
    
    # Return the created collaborator with its ID
    collaborator["_id"] = str(result.inserted_id)
    
    print(f"Collaborator {email} added to project {project_id} by {username}")
    
    return jsonify({
        "message": "Collaborator added successfully",
        "collaborator": collaborator
    })

# Enhanced remove collaborator endpoint with better error handling
@app.route("/collaborators/<collaborator_id>", methods=["DELETE"])
@login_required
def remove_collaborator(collaborator_id):
    username = session["user"]
    
    # Get user details to check role
    user = users_collection.find_one({"username": username})
    user_role = user.get("role", "user") if user else "user"
    
    try:
        object_id = ObjectId(collaborator_id)
    except:
        return jsonify({"error": "Invalid collaborator ID"}), 400
    
    # Find the collaborator
    collaborator = collaborators_collection.find_one({"_id": object_id})
    if not collaborator:
        return jsonify({"error": "Collaborator not found"}), 404
    
    # Find the project
    project = projects_collection.find_one({"id": collaborator["project_id"]})
    if not project:
        return jsonify({"error": "Project not found"}), 404
    
    # Check permissions: only owner or admin can remove collaborators
    if project["user"] != username and user_role != "admin":
        return jsonify({"error": "Only project owners and admins can remove collaborators"}), 403
    
    # Remove from collaborators collection
    collaborators_collection.delete_one({"_id": object_id})
    
    # Remove from project collaborators list
    projects_collection.update_one(
        {"id": collaborator["project_id"]},
        {"$pull": {"collaborators": collaborator["username"]}}
    )
    
    print(f"Collaborator {collaborator['email']} removed from project {collaborator['project_id']} by {username}")
    
    return jsonify({"message": "Collaborator removed successfully"})

# File export endpoints
@app.route("/download_requirements/<format>", methods=["POST"])
@login_required
def download_requirements(format):
    """Generate a file with the requirements in the specified format (PDF or DOCX)"""
    data = request.json
    requirements = data.get("requirements", [])
    project_name = data.get("project_name", "Projet")
    format_date = datetime.now().strftime("%Y-%m-%d")
    file_name = f"exigences_{project_name.replace(' ', '_')}_{format_date}"
    
    if not requirements:
        return jsonify({"error": "No requirements provided"}), 400
    
    username = session["user"]
    
    if format == "pdf":
        # Generate PDF using ReportLab
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=42,
            leftMargin=42,
            topMargin=42,
            bottomMargin=42
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkblue
        )
        
        normal_style = styles['Normal']
        
        elements = []
        
        # Add title
        title = Paragraph(f"Exigences: {project_name}", title_style)
        elements.append(title)
        
        # Add date
        date_para = Paragraph(f"Date: {format_date}", normal_style)
        elements.append(date_para)
        elements.append(Spacer(1, 0.3*inch))
        
        for req in requirements:
            # Requirement title
            req_title = Paragraph(req.get("title", ""), heading_style)
            elements.append(req_title)
            
            # Requirement details in table format
            req_details = [
                f"Catégorie: {get_category_label(req.get('category', ''))}",
                f"Priorité: {get_priority_label(req.get('priority', ''))}",
                f"Statut: {get_status_label(req.get('status', ''))}"
            ]
            
            for detail in req_details:
                detail_para = Paragraph(detail, normal_style)
                elements.append(detail_para)
            
            # Description
            desc_title = Paragraph("Description:", normal_style)
            elements.append(desc_title)
            desc_para = Paragraph(req.get("description", ""), normal_style)
            elements.append(desc_para)
            
            elements.append(Spacer(1, 0.2*inch))
        
        # Build and save the PDF
        doc.build(elements)
        
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{file_name}.pdf",
            mimetype="application/pdf"
        )
    
    elif format == "docx":
        # Generate DOCX using python-docx
        doc = Document()
        
        # Add document title
        title = doc.add_heading(f"Exigences: {project_name}", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph(f"Date: {format_date}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Add spacing
        
        for req in requirements:
            # Requirement title
            doc.add_heading(req.get("title", ""), level=1)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.autofit = True
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "Métadonnée"
            header_cells[1].text = "Valeur"
            
            row_cells = table.add_row().cells
            row_cells[0].text = "Catégorie"
            row_cells[1].text = get_category_label(req.get("category", ""))
            
            row_cells = table.add_row().cells
            row_cells[0].text = "Priorité"
            row_cells[1].text = get_priority_label(req.get("priority", ""))
            
            row_cells = table.add_row().cells
            row_cells[0].text = "Statut"
            row_cells[1].text = get_status_label(req.get("status", ""))
            
            doc.add_paragraph().add_run("Description:").bold = True
            doc.add_paragraph(req.get("description", ""))
            
            doc.add_paragraph()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{file_name}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    else:
        return jsonify({"error": f"Unsupported format: {format}"}), 400

@app.route("/export_test_cases", methods=["POST"])
@login_required
def export_test_cases():
    """Generate a file with test cases in the specified format (PDF or DOCX)"""
    data = request.json
    test_cases = data.get("test_cases", "")
    format_type = data.get("format", "pdf")  # 'pdf' or 'docx'
    project_id = data.get("project_id", "")
    requirement_id = data.get("requirement_id", "")
    requirement_title = data.get("requirement_title", "Test Cases")
    
    if not test_cases:
        return jsonify({"error": "No test cases provided"}), 400
    
    username = session["user"]
    format_date = datetime.now().strftime("%Y-%m-%d")
    
    # Sanitize filename
    safe_title = requirement_title.replace(" ", "_").replace("/", "_").replace("\\", "_")
    file_name = f"test_cases_{safe_title}_{format_date}"
    
    # Verify project access if project_id is provided
    if project_id:
        project = projects_collection.find_one({
            "id": project_id,
            "$or": [
                {"user": username},
                {"collaborators": username}
            ]
        })
        
        if not project:
            return jsonify({"error": "Project not found or access denied"}), 403
    
    # Process based on format type
    if format_type == "pdf":
        try:
            # Generate PDF using ReportLab
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=42,
                leftMargin=42,
                topMargin=42,
                bottomMargin=42
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1
            )
            
            normal_style = styles['Normal']
            
            elements = []
            
            # Add title
            title = Paragraph(f"Test Cases: {requirement_title}", title_style)
            elements.append(title)
            
            # Add date
            date_para = Paragraph(f"Date: {format_date}", normal_style)
            elements.append(date_para)
            elements.append(Spacer(1, 0.3*inch))
            
            # Add test cases with monospaced font to preserve formatting
            test_case_style = ParagraphStyle(
                'TestCaseStyle',
                parent=normal_style,
                fontName='Courier',
                fontSize=9,
                leftIndent=0,
                rightIndent=0
            )
            
            # Split test cases into lines and add each as a paragraph to preserve formatting
            for line in test_cases.split('\n'):
                if line.strip():
                    para = Paragraph(line, test_case_style)
                    elements.append(para)
                else:
                    elements.append(Spacer(1, 0.1*inch))
            
            # Build and save the PDF
            doc.build(elements)
            
            buffer.seek(0)
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f"{file_name}.pdf",
                mimetype="application/pdf"
            )
            
        except Exception as e:
            print(f"Error generating PDF: {str(e)}")
            return jsonify({"error": f"Failed to generate PDF: {str(e)}"}), 500
            
    elif format_type == "docx":
        try:
            # Generate DOCX using python-docx
            doc = Document()
            
            # Add document title
            title = doc.add_heading(f"Test Cases: {requirement_title}", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date
            date_paragraph = doc.add_paragraph(f"Date: {format_date}")
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.add_paragraph()  # Add spacing
            
            # Add test cases with monospaced font to preserve formatting
            test_case_para = doc.add_paragraph()
            test_case_run = test_case_para.add_run(test_cases)
            test_case_run.font.name = 'Courier New'
            test_case_run.font.size = Pt(9)
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name=f"{file_name}.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            print(f"Error generating DOCX: {str(e)}")
            return jsonify({"error": f"Failed to generate DOCX: {str(e)}"}), 500
    
    else:
        return jsonify({"error": f"Unsupported format: {format_type}"}), 400




def get_effective_llm_service(username, project_id=None):
    """Determine which LLM service to use - ONLY based on project settings"""
    try:
        if project_id:
            # Use project-level settings ONLY
            return get_effective_llm_for_project(project_id)
        else:
            # No project context, default to local RAG if available
            return "minio" if check_rag_availability() else None
                
    except Exception as e:
        print(f"Error determining effective LLM service: {e}")
        return "minio" if check_rag_availability() else None




def validate_claude_api_key(api_key):
    """Validate Claude API key by making a test request"""
    if not api_key or not api_key.strip():
        return False, "API key is required"
    
    try:
        clean_key = api_key.strip()
        print(f"Validating Claude API key: {clean_key[:20]}...")
        
        if not (clean_key.startswith("sk-ant-") or clean_key.startswith("sk-")):
            return False, "Invalid Claude API key format (should start with 'sk-ant-' or 'sk-')"
        
        client = anthropic.Anthropic(api_key=clean_key)
        
        response = client.messages.create(
            model="claude-3-haiku-20240307",
            max_tokens=10,
            messages=[{"role": "user", "content": "Hi"}]
        )
        
        print("Claude API key validation successful")
        return True, "API key is valid"
    
    except anthropic.AuthenticationError as e:
        print(f"Claude authentication error: {e}")
        return False, "Invalid API key or authentication failed"
    except anthropic.PermissionDeniedError as e:
        print(f"Claude permission error: {e}")
        return False, "API key doesn't have required permissions"
    except anthropic.RateLimitError as e:
        print(f"Claude rate limit error: {e}")
        return False, "Rate limit exceeded, but API key appears valid"
    except Exception as e:
        print(f"Claude API key validation error: {e}")
        return False, f"API key validation failed: {str(e)}"



def save_project_llm_settings(project_id, manager_username, llm_choice, api_key=None):
    """Save LLM settings at project level - CLAUDE AND MINIO ONLY"""
    try:
        print(f"Saving project LLM settings for project {project_id} by {manager_username}")
        
        # Verify manager permissions
        user = users_collection.find_one({"username": manager_username})
        if not user or user.get("role") not in ["manager", "admin"]:
            print(f"Access denied: User {manager_username} role is {user.get('role') if user else 'None'}")
            return False, "Only managers can configure project LLM settings"
        
        # Verify manager owns the project
        project = projects_collection.find_one({"id": project_id})
        if not project:
            print(f"Project not found: {project_id}")
            return False, "Project not found"
        
        if user.get("role") != "admin" and project["user"] != manager_username:
            print(f"Permission denied: Project owner is {project['user']}, requesting user is {manager_username}")
            return False, "You can only configure settings for projects you manage"
        
        # Prepare settings data
        settings_data = {
            "project_id": project_id,
            "manager": manager_username,
            "llm_service": llm_choice,
            "updated_at": datetime.now(timezone.utc)
        }
        
        # Handle API key validation for Claude
        if llm_choice == "claude":
            if not api_key:
                print("No API key provided for Claude service")
                return False, "API key is required for Claude"
            
            print(f"Validating Claude API key for project {project_id}")
            # Validate API key
            is_valid, message = validate_claude_api_key(api_key)
            if not is_valid:
                print(f"API key validation failed: {message}")
                return False, f"API key validation failed: {message}"
            
            settings_data["claude_api_key"] = api_key.strip()
            settings_data["api_key_validated"] = True
            settings_data["api_key_validated_at"] = datetime.now(timezone.utc)
            print("API key validated and stored successfully")
            
        elif llm_choice == "minio":
            # Remove API key if switching to MinIO
            settings_data["claude_api_key"] = None
            settings_data["api_key_validated"] = False
            print("Switched to MinIO service, removed API key")
        
        # Check if settings already exist
        existing_settings = project_settings_collection.find_one({"project_id": project_id})
        
        if existing_settings:
            # Update existing settings
            settings_data["created_at"] = existing_settings.get("created_at", datetime.now(timezone.utc))
            project_settings_collection.update_one(
                {"project_id": project_id},
                {"$set": settings_data}
            )
            print("Updated existing project settings")
        else:
            # Create new settings
            settings_data["created_at"] = datetime.now(timezone.utc)
            project_settings_collection.insert_one(settings_data)
            print("Created new project settings")
        
        return True, "Project LLM settings saved successfully"
        
    except Exception as e:
        print(f"Error saving project LLM settings: {e}")
        import traceback
        traceback.print_exc()
        return False, f"Error saving settings: {str(e)}"


def get_user_llm_settings(username, project_id=None):
    """Get user's LLM settings for a specific project or global"""
    try:
        # Try to get project-specific settings first if project_id is provided
        if project_id:
            project_settings = user_settings_collection.find_one({
                "user": username,
                "project_id": project_id
            })
            if project_settings:
                return project_settings
        
        # Fallback to global settings
        global_settings = user_settings_collection.find_one({
            "user": username,
            "project_id": {"$exists": False}
        })
        
        return global_settings
        
    except Exception as e:
        print(f"Error getting user LLM settings: {e}")
        return None

def get_effective_llm_for_user(username, project_id=None):
    """Determine which LLM service the user should use - CLAUDE OR LOCAL ONLY"""
    try:
        settings = get_user_llm_settings(username, project_id)
        
        if not settings:
            return "minio" if check_rag_availability() else None
        
        llm_service = settings.get("llm_service", "local")
        
        if llm_service == "claude":
            if settings.get("claude_api_key") and settings.get("api_key_validated"):
                return "claude"
            else:
                return "minio" if check_rag_availability() else None
        else:
            return "minio" if check_rag_availability() else None
            
    except Exception as e:
        print(f"Error determining effective LLM service: {e}")
        return "minio" if check_rag_availability() else None




def save_project_llm_settings(project_id, manager_username, llm_choice, api_key=None):
    """Save LLM settings at project level - enhanced with debugging"""
    try:
        print(f"Saving project LLM settings for project {project_id} by {manager_username}")
        
        # Verify manager permissions
        user = users_collection.find_one({"username": manager_username})
        if not user or user.get("role") not in ["manager", "admin"]:
            print(f"Access denied: User {manager_username} role is {user.get('role') if user else 'None'}")
            return False, "Only managers can configure project LLM settings"
        
        # Verify manager owns the project
        project = projects_collection.find_one({"id": project_id})
        if not project:
            print(f"Project not found: {project_id}")
            return False, "Project not found"
        
        if user.get("role") != "admin" and project["user"] != manager_username:
            print(f"Permission denied: Project owner is {project['user']}, requesting user is {manager_username}")
            return False, "You can only configure settings for projects you manage"
        
        # Prepare settings data
        settings_data = {
            "project_id": project_id,
            "manager": manager_username,
            "llm_service": llm_choice,
            "updated_at": datetime.now(timezone.utc)
        }
        
        # Handle API key validation for Claude
        if llm_choice == "claude":
            if not api_key:
                print("No API key provided for Claude service")
                return False, "API key is required for Claude"
            
            print(f"Validating Claude API key for project {project_id}")
            # Validate API key
            is_valid, message = validate_claude_api_key(api_key)
            if not is_valid:
                print(f"API key validation failed: {message}")
                return False, f"API key validation failed: {message}"
            
            settings_data["claude_api_key"] = api_key.strip()
            settings_data["api_key_validated"] = True
            settings_data["api_key_validated_at"] = datetime.now(timezone.utc)
            print("API key validated and stored successfully")
            
        elif llm_choice == "minio":
            # Remove API key if switching to local
            settings_data["claude_api_key"] = None
            settings_data["api_key_validated"] = False
            print("Switched to local service, removed API key")
        
        # Check if settings already exist
        existing_settings = project_settings_collection.find_one({"project_id": project_id})
        
        if existing_settings:
            # Update existing settings
            settings_data["created_at"] = existing_settings.get("created_at", datetime.now(timezone.utc))
            project_settings_collection.update_one(
                {"project_id": project_id},
                {"$set": settings_data}
            )
            print("Updated existing project settings")
        else:
            # Create new settings
            settings_data["created_at"] = datetime.now(timezone.utc)
            project_settings_collection.insert_one(settings_data)
            print("Created new project settings")
        
        return True, "Project LLM settings saved successfully"
        
    except Exception as e:
        print(f"Error saving project LLM settings: {e}")
        import traceback
        traceback.print_exc()
        return False, f"Error saving settings: {str(e)}"
def get_project_llm_settings(project_id):
    """Get LLM settings for a specific project"""
    try:
        return project_settings_collection.find_one({"project_id": project_id})
    except Exception as e:
        print(f"Error getting project LLM settings: {e}")
        return None

def get_effective_llm_for_project(project_id):
    """Determine which LLM service should be used for a project - CLAUDE OR MINIO ONLY"""
    try:
        settings = get_project_llm_settings(project_id)
        
        if not settings:
            # No project settings found, default to minio if available
            return "minio" if check_rag_availability() else None
        
        llm_service = settings.get("llm_service", "minio")
        
        if llm_service == "claude":
            # Check if API key is available and validated
            if settings.get("claude_api_key") and settings.get("api_key_validated"):
                return "claude"
            else:
                # API key not available or not validated, fallback to minio
                return "minio" if check_rag_availability() else None
        else:
            # Project configured for minio
            return "minio" if check_rag_availability() else None
            
    except Exception as e:
        print(f"Error determining effective LLM service for project: {e}")
        return "minio" if check_rag_availability() else None
def get_project_api_key(project_id):
    """Get the API key for a project"""
    try:
        settings = get_project_llm_settings(project_id)
        if settings and settings.get("claude_api_key") and settings.get("api_key_validated"):
            return settings["claude_api_key"]
        return None
    except Exception as e:
        print(f"Error getting project API key: {e}")
        return None
@app.route("/check_api_service", methods=["GET"])
@login_required
def check_api_service():
    """Check which AI service is configured for the current context - CLAUDE OR LOCAL ONLY"""
    try:
        project_id = request.args.get("project_id")
        
        if project_id:
            # Check project-level configuration
            effective_service = get_effective_llm_for_project(project_id)
            
            if effective_service == "claude":
                return jsonify({
                    "configured": True,
                    "service": "claude",
                    "service_name": "Claude AI",
                    "message": "Using Claude AI for this project",
                    "project_configured": True
                })
            elif effective_service == "minio":
                    return jsonify({
                        "service": "minio", 
                        "service_name": "MinIO RAG",
                        "message": "Using MinIO RAG for this project", 
                    "project_configured": True,
                    "rag_available": check_rag_availability()
                })
            else:
                return jsonify({
                    "configured": False,
                    "service": None,
                    "message": "No AI service configured for this project. Please configure Claude API or ensure Local RAG is available.",
                    "project_configured": False
                })
        else:
            # No project context - require project configuration
            return jsonify({
                "configured": False,
                "service": None,
                "message": "Project ID required for AI service configuration"
            })
        
    except Exception as e:
        print(f"Error checking API service: {e}")
        return jsonify({"error": str(e)}), 500
    
@app.route("/check_global_api_service", methods=["GET"])
@login_required
def check_global_api_service():
    """Check if any global API service is configured (for fallback)"""
    try:
        # Since you're only using project-level settings now, 
        # this endpoint can just return local RAG availability
        return jsonify({
            "configured": check_rag_availability(),
            "service": "minio" if check_rag_availability() else None,
            "service_name": "MinIO RAG" if check_rag_availability() else None
        })
        
    except Exception as e:
        print(f"Error checking global API service: {e}")
        return jsonify({"error": str(e)}), 500
@app.route("/project_llm_settings/<project_id>", methods=["GET"])
@login_required
def get_project_llm_settings_endpoint(project_id):
    username = session["user"]
    
    try:
        user = users_collection.find_one({"username": username})
        if not user:
            return jsonify({"error": "User not found"}), 404
        
        project = projects_collection.find_one({"id": project_id})
        if not project:
            return jsonify({"error": "Project not found"}), 404
        
        user_role = user.get("role", "user")
        
        has_access = (
            project["user"] == username or
            username in project.get("collaborators", []) or
            username in project.get("assigned_users", []) or
            user_role in ["manager", "admin"]
        )
        
        if not has_access:
            return jsonify({"error": "Access denied to this project"}), 403
        
        can_modify = user_role in ["manager", "admin"] and (
            project["user"] == username or user_role == "admin"
        )
        
        settings = get_project_llm_settings(project_id)
        
        response_data = {
            "project_id": project_id,
            "llm_service": settings.get("llm_service", "minio") if settings else "minio",
            "has_claude_key": bool(settings and settings.get("claude_api_key")) if can_modify else None,
            "api_key_validated": settings.get("api_key_validated", False) if settings else False,
            "manager": settings.get("manager") if settings else None,
            "minio_available": check_rag_availability(),  # Changé de 'local_available' à 'minio_available'
            "effective_service": get_effective_llm_for_project(project_id),
            "can_modify_settings": can_modify,
            "user_role": user_role
        }
        
        return jsonify(response_data)
        
    except Exception as e:
        print(f"Error getting project LLM settings: {e}")
        return jsonify({"error": str(e)}), 500
@app.route("/validate_claude_key", methods=["POST"])
@login_required
def validate_claude_key_endpoint():
    """Endpoint to validate Claude API key - called by frontend"""
    try:
        data = request.json
        api_key = data.get("api_key", "").strip()
        
        if not api_key:
            return jsonify({"valid": False, "message": "API key is required"}), 400
        
        print(f"Validating Claude API key: {api_key[:20]}...")
        
        # Use the corrected validation function
        is_valid, message = validate_claude_api_key(api_key)
        
        if is_valid:
            print("Claude API key validation successful")
            return jsonify({"valid": True, "message": message})
        else:
            print(f"Claude API key validation failed: {message}")
            return jsonify({"valid": False, "message": message}), 400
            
    except Exception as e:
        print(f"Error in validate_claude_key_endpoint: {e}")
        return jsonify({"valid": False, "message": f"Validation error: {str(e)}"}), 500

@app.route("/project_llm_settings/<project_id>", methods=["POST"])
@login_required
def save_project_llm_settings_endpoint(project_id):
    """Save project's LLM settings - MinIO and Claude only"""
    username = session["user"]
    data = request.json
    
    print(f"Received project LLM settings request for project {project_id} from user {username}")
    print(f"Request data: {data}")
    
    llm_choice = data.get("llm_service")
    api_key = data.get("claude_api_key", "").strip() if data.get("claude_api_key") else None
    
    print(f"llm_choice: {llm_choice}")
    print(f"api_key provided: {'Yes' if api_key else 'No'}")
    
    # Validate input - ONLY MINIO OR CLAUDE
    if llm_choice not in ["minio", "claude"]:
        print(f"Invalid llm_service: {llm_choice}")
        return jsonify({"error": "llm_service must be 'minio' or 'claude'"}), 400
    
    if llm_choice == "claude" and not api_key:
        print("No API key provided for Claude service")
        return jsonify({"error": "API key is required when selecting Claude"}), 400
    
    try:
        success, message = save_project_llm_settings(project_id, username, llm_choice, api_key)
        
        if success:
            print(f"Successfully saved settings: {message}")
            updated_settings = get_project_llm_settings(project_id)
            
            response_data = {
                "message": message,
                "settings": {
                    "project_id": project_id,
                    "llm_service": updated_settings.get("llm_service"),
                    "has_claude_key": bool(updated_settings.get("claude_api_key")),
                    "api_key_validated": updated_settings.get("api_key_validated", False),
                    "effective_service": get_effective_llm_for_project(project_id)
                }
            }
            
            return jsonify(response_data)
        else:
            print(f"Failed to save settings: {message}")
            return jsonify({"error": message}), 400
            
    except Exception as e:
        print(f"Error in save_project_llm_settings_endpoint: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
@app.route("/project_llm_settings/<project_id>", methods=["DELETE"])
@login_required
def delete_project_llm_settings_endpoint(project_id):
    """Delete project LLM settings (reset to defaults)"""
    username = session["user"]
    
    try:
        user = users_collection.find_one({"username": username})
        if not user or user.get("role") not in ["manager", "admin"]:
            return jsonify({"error": "Access denied"}), 403
        
        project = projects_collection.find_one({"id": project_id})
        if not project:
            return jsonify({"error": "Project not found"}), 404
        
        if user.get("role") != "admin" and project["user"] != username:
            return jsonify({"error": "You can only delete settings for projects you manage"}), 403
        
        # Delete the settings
        project_settings_collection.delete_one({"project_id": project_id})
        
        return jsonify({"message": "Project LLM settings reset to defaults"})
        
    except Exception as e:
        print(f"Error deleting project LLM settings: {e}")
        return jsonify({"error": str(e)}), 500
# Add this to your app.py

@app.route("/project_info/<project_id>", methods=["GET"])
@login_required
def get_project_info_for_users(project_id):
    """Get project information for regular users (read-only)"""
    username = session["user"]
    
    try:
        # Verify user has access to this project
        project = projects_collection.find_one({"id": project_id})
        if not project:
            return jsonify({"error": "Project not found"}), 404
        
        # Check if user is owner, collaborator, or assigned user
        user = users_collection.find_one({"username": username})
        user_role = user.get("role", "user")
        
        has_access = (
            project["user"] == username or  # Owner
            username in project.get("collaborators", []) or  # Collaborator
            username in project.get("assigned_users", []) or  # Assigned user
            user_role in ["manager", "admin"]  # Manager/Admin
        )
        
        if not has_access:
            return jsonify({"error": "Access denied to this project"}), 403
        
        # Get project settings (what model is configured)
        project_settings = get_project_llm_settings(project_id)
        effective_service = get_effective_llm_for_project(project_id)
        
        # Prepare response with safe information
        project_info = {
            "project_id": project_id,
            "name": project.get("name"),
            "effective_service": effective_service,
            "service_display_name": {
                "minio": "MinIO RAG System",
                "claude": "Claude AI"
            }.get(effective_service, effective_service),
            "is_configured": bool(project_settings),
            "manager": project_settings.get("manager") if project_settings else project.get("user"),
            "user_role": user_role,
            "can_modify_settings": user_role in ["manager", "admin"] and (
                project["user"] == username or user_role == "admin"
            )
        }
        
        return jsonify(project_info)
        
    except Exception as e:
        print(f"Error getting project info: {e}")
        return jsonify({"error": str(e)}), 500
@app.after_request
def after_request(response):
    return response

if not initialize_minio_rag():
    print("❌ CRITICAL ERROR: Cannot start application without MinIO RAG System")
    print("❌ Please ensure MinIO is running and minio_rag_system.py is available")
    raise RuntimeError("MinIO RAG initialization failed")

if __name__ == "__main__":
    app.run(debug=False, host="0.0.0.0", port=5000, use_reloader=False, threaded=True)